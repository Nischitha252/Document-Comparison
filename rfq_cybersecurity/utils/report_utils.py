"""
Report generation utilities.

This module provides helper functions for converting the nested
result dictionary into a Word document and for encoding images as
data URIs suitable for embedding in HTML/CSS used by Streamlit.
Placing these functions here separates pure formatting concerns from
the core business logic.
"""

from __future__ import annotations

import base64
import os
from typing import Dict

from docx import Document as DocxDocument

# Import the package logger
from ..logging_config import get_logger

logger = get_logger(__name__)


def json_to_docx(result: Dict[str, Dict[str, str]]) -> DocxDocument:
    """Convert the nested result dictionary into a Word document.

    A document is created with a top‑level title followed by
    headings corresponding to each cybersecurity topic.  Within
    each topic, the relevant sections are added as sub‑headings
    along with the summarised text.

    Parameters
    ----------
    result: dict
        The processed results mapping heading names to sections and
        summaries.

    Returns
    -------
    docx.Document
        A populated Word document object.
    """
    # Create a new Word document
    doc = DocxDocument()
    # Add a top‑level title
    doc.add_heading("Cybersecurity Report", level=0)
    # Iterate over each heading in the result
    for heading, sections in result.items():
        # Add a second‑level heading for the topic
        doc.add_heading(heading, level=1)
        # Within each topic, add sub‑headings and their summaries
        for section_name, snippet in sections.items():
            doc.add_heading(section_name, level=2)
            doc.add_paragraph(snippet)
    logger.debug("Generated Word document for %d headings", len(result))
    return doc


def load_header_image(image_path: str) -> str:
    """Encode an image file as a base64 data URI.

    Streamlit allows us to embed images via HTML.  This helper reads
    the given image from disk and returns a base64‑encoded string
    suitable for embedding into CSS backgrounds or inline HTML.

    Parameters
    ----------
    image_path: str
        Absolute path to the image file.

    Returns
    -------
    str
        Base64 encoded representation of the image prefixed with
        ``data:image/...;base64,``.
    """
    try:
        # Open the image file in binary mode and read its contents
        with open(image_path, "rb") as img_file:
            encoded_bytes = base64.b64encode(img_file.read())
    except Exception as exc:
        logger.error("Could not load header image '%s': %s", image_path, exc)
        # Return an empty data URI on failure to avoid breaking the UI
        return ""
    # Derive the file extension to construct the correct MIME type
    extension = os.path.splitext(image_path)[1].lstrip(".").lower()
    data_uri = f"data:image/{extension};base64,{encoded_bytes.decode()}"
    logger.debug("Loaded header image '%s' (extension=%s)", image_path, extension)
    return data_uri